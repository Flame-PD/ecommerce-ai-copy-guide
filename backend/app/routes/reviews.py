from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from sqlalchemy import desc, func
from typing import List, Optional
import io
import re
import pandas as pd
from docx import Document
from fastapi.responses import StreamingResponse
from app.database import get_db
from app.models.user import User
from app.models.product import Product
from app.models.review import Review
from app.models.order import Order
from app.schemas import ReviewCreate, ReviewOut, SentimentStats, MerchantReviewCreate
from app.utils.security import get_current_user, require_merchant
from app.services import sentiment_service, ai_service
from app.utils.helpers import save_upload_file

router = APIRouter()


@router.post("", response_model=ReviewOut)
def create_review(
    product_id: int = Form(...),
    order_id: Optional[int] = Form(None),
    rating: int = Form(...),
    content: str = Form(...),
    images: List[UploadFile] = File(default=[]),
    video: Optional[UploadFile] = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    product = db.query(Product).filter(Product.id == product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="商品不存在")
    if order_id:
        order = db.query(Order).filter(Order.id == order_id, Order.user_id == current_user.id).first()
        if not order:
            raise HTTPException(status_code=403, detail="无权评价该订单")
    sentiment = sentiment_service.analyze(content)
    image_urls = []
    for img in images:
        if img.filename:
            image_urls.append(save_upload_file(img, "reviews/images"))
    video_url = ""
    if video and video.filename:
        video_url = save_upload_file(video, "reviews/videos")
    review = Review(
        user_id=current_user.id,
        product_id=product_id,
        order_id=order_id,
        rating=rating,
        content=content,
        sentiment=sentiment,
        source="user",
        images=image_urls,
        video_url=video_url,
    )
    db.add(review)
    db.commit()
    db.refresh(review)
    return review


@router.post("/merchant-create", response_model=ReviewOut)
def merchant_create_review(payload: MerchantReviewCreate, db: Session = Depends(get_db), current_user: User = Depends(require_merchant)):
    product = db.query(Product).filter(Product.id == payload.product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="商品不存在")
    sentiment = sentiment_service.analyze(payload.content)
    review = Review(
        user_id=current_user.id,
        product_id=payload.product_id,
        rating=payload.rating,
        content=payload.content,
        sentiment=sentiment,
        source="merchant",
    )
    db.add(review)
    db.commit()
    db.refresh(review)
    return review


@router.get("/product/{product_id}", response_model=List[ReviewOut])
def product_reviews(product_id: int, sentiment: str = None, db: Session = Depends(get_db)):
    q = db.query(Review).filter(Review.product_id == product_id, Review.source == "user")
    if sentiment:
        q = q.filter(Review.sentiment == sentiment)
    return q.order_by(desc(Review.created_at)).all()


@router.get("/merchant/list")
def merchant_reviews(
    product_id: int = None,
    keyword: str = None,
    sentiment: str = None,
    status: str = None,
    db: Session = Depends(get_db),
    _: User = Depends(require_merchant),
):
    q = db.query(Review, Product, User).join(Product, Review.product_id == Product.id).join(User, Review.user_id == User.id)
    if product_id:
        q = q.filter(Review.product_id == product_id)
    if keyword:
        q = q.filter(Review.content.contains(keyword))
    if sentiment:
        q = q.filter(Review.sentiment == sentiment)
    if status:
        q = q.filter(Product.status == status)
    rows = q.order_by(desc(Review.created_at)).all()
    return [
        {
            "id": r.id,
            "user_id": r.user_id,
            "product_id": r.product_id,
            "order_id": r.order_id,
            "rating": r.rating,
            "content": r.content,
            "sentiment": r.sentiment,
            "source": r.source,
            "images": r.images or [],
            "video_url": r.video_url or "",
            "created_at": r.created_at,
            "product_name": p.name,
            "product_display_id": p.display_id,
            "product_status": p.status,
            "user_nickname": u.nickname,
            "user_display_id": u.display_id,
        }
        for r, p, u in rows
    ]


@router.get("/stats/{product_id}", response_model=SentimentStats)
def review_stats(product_id: int, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    reviews = db.query(Review).filter(Review.product_id == product_id).all()
    result = sentiment_service.batch_analyze([{"content": r.content, "rating": r.rating} for r in reviews])
    return SentimentStats(**result)


@router.post("/summary")
def review_summary(payload: dict, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    product_id = payload.get("product_id")
    reviews = db.query(Review).filter(Review.product_id == product_id).all()
    data = [{"content": r.content, "rating": r.rating, "sentiment": r.sentiment} for r in reviews]
    result = ai_service.generate_review_summary(data)
    return {"ok": True, **result}


@router.post("/summary/export")
def export_summary(payload: dict, _: User = Depends(require_merchant)):
    fmt = payload.get("format", "txt")
    content = payload.get("content", "")
    title = payload.get("title", "评论总结")
    if fmt == "docx":
        doc = Document()
        doc.add_heading(title, level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={title}.docx"})
    else:
        buffer = io.BytesIO(content.encode("utf-8"))
        return StreamingResponse(buffer, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename={title}.txt"})


@router.get("/template/download")
def download_template(_: User = Depends(require_merchant)):
    df = pd.DataFrame([{
        "product_name": "商品名称",
        "rating": 5,
        "content": "评论内容",
    }])
    buffer = io.BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers={"Content-Disposition": "attachment; filename=reviews_template.xlsx"})


@router.post("/analyze")
def analyze_upload(
    file: UploadFile = File(...),
    product_id: int = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_merchant),
):
    ext = file.filename.split(".")[-1].lower()
    content = file.file.read()
    texts = []
    if ext == "txt" or ext == "md":
        texts = [t.strip() for t in content.decode("utf-8", errors="ignore").split("\n") if t.strip()]
    elif ext in ("xlsx", "xls"):
        df = pd.read_excel(io.BytesIO(content))
        col = "content" if "content" in df.columns else df.columns[0]
        texts = [str(t).strip() for t in df[col].dropna().tolist()]
    elif ext == "docx":
        doc = Document(io.BytesIO(content))
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    else:
        raise HTTPException(status_code=400, detail="不支持的文件格式")

    reviews = [{"content": t, "rating": 5} for t in texts]
    result = sentiment_service.batch_analyze(reviews)
    if product_id:
        for t in texts:
            sentiment = sentiment_service.analyze(t)
            db.add(Review(user_id=current_user.id, product_id=product_id, rating=5, content=t, sentiment=sentiment, source="merchant"))
        db.commit()
    return {"ok": True, "count": len(texts), "stats": result}


@router.post("/import")
def import_reviews(file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(require_merchant)):
    ext = file.filename.split(".")[-1].lower()
    if ext not in ("xlsx", "xls"):
        raise HTTPException(status_code=400, detail="仅支持xlsx/xls")
    df = pd.read_excel(io.BytesIO(file.file.read()))
    rows = df.to_dict(orient="records")
    products = {p.name: p.id for p in db.query(Product).all()}
    count = 0
    for row in rows:
        product_name = str(row.get("product_name", "") or "").strip()
        product_id = products.get(product_name)
        if not product_id:
            continue
        rating = int(row.get("rating", 5) or 5)
        content = str(row.get("content", "") or "").strip()
        if not content:
            continue
        sentiment = sentiment_service.analyze(content)
        db.add(Review(user_id=current_user.id, product_id=product_id, rating=rating, content=content, sentiment=sentiment, source="merchant"))
        count += 1
    db.commit()
    return {"ok": True, "count": count}


@router.delete("/{review_id}")
def delete_review(review_id: int, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    review = db.query(Review).filter(Review.id == review_id).first()
    if not review:
        raise HTTPException(status_code=404, detail="评论不存在")
    db.delete(review)
    db.commit()
    return {"ok": True}
