from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.user import User
from app.models.product import Product
from app.models.live_script import LiveScript
from app.schemas import AICopyRequest, AIScriptRequest, AIResponse
from app.utils.security import require_merchant
from app.services import ai_service
from app.services.ai_client import get_ai_client
import logging

logger = logging.getLogger(__name__)
router = APIRouter()


def _product_to_dict(product: Product) -> dict:
    """Convert a Product ORM object to a plain dict for AI consumption."""
    return {
        "id": product.id,
        "name": product.name,
        "category": product.category or "",
        "description": product.description or "",
        "price": product.price or 0.0,
        "specs": product.specs or [],
    }


@router.post("/copy", response_model=AIResponse)
def generate_copy(payload: AICopyRequest, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    # -- resolve product data --------------------------------------------------
    if payload.product_id:
        product = db.query(Product).filter(Product.id == payload.product_id).first()
        if not product:
            raise HTTPException(status_code=404, detail="商品不存在")
        data = _product_to_dict(product)
    else:
        if not payload.name:
            raise HTTPException(status_code=400, detail="请输入商品名称")
        data = {
            "id": None,
            "name": payload.name,
            "category": payload.category or "",
            "description": payload.description or "",
            "price": payload.price or 0.0,
            "specs": payload.specs or [],
        }

    # -- try AI service first, fall back to local ------------------------------
    ai = get_ai_client()
    features = ", ".join(data.get("specs", [])) if data.get("specs") else data.get("description", "")

    # 1) Title
    titles = ai.generate_titles(data["name"], features, data["category"])
    # 2) Description
    desc = ai.generate_description(data["name"], features, data["category"],
                                    specifications=", ".join(data.get("specs", [])),
                                    target_audience="通用")

    if titles is not None and desc is not None:
        # AI service responded successfully
        result = {
            "title": titles[0] if titles else "",
            "selling_points": "\n".join(f"• {t}" for t in titles[1:]) if len(titles) > 1 else "",
            "detail": desc,
            "slogan": titles[-1] if titles else "",
        }
        logger.info("AI copy generated via AI service for '%s'", data["name"])
    else:
        # Fall back to local ai_service (direct DeepSeek)
        logger.warning("AI service unavailable — falling back to local ai_service")
        result = ai_service.generate_product_copy(data, payload.style)

    # -- persist ----------------------------------------------------------------
    if payload.product_id:
        product = db.query(Product).filter(Product.id == payload.product_id).first()
        if product:
            product.ai_title = result.get("title", "")
            product.ai_selling_points = result.get("selling_points", "")
            product.ai_detail = result.get("detail", "")
            product.ai_slogan = result.get("slogan", "")
            db.commit()
    return AIResponse(result=result)


@router.post("/script", response_model=AIResponse)
def generate_script(payload: AIScriptRequest, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    product = db.query(Product).filter(Product.id == payload.product_id).first()
    if not product:
        raise HTTPException(status_code=404, detail="商品不存在")
    data = _product_to_dict(product)

    # -- try AI service first ---------------------------------------------------
    ai = get_ai_client()
    features = ", ".join(data.get("specs", [])) if data.get("specs") else data.get("description", "")
    live_data = ai.generate_livestream(
        data["name"], features, data["category"],
        promotion="", target_audience="通用",
    )

    if live_data is not None:
        content = live_data.get("script", "")
        title = f"{product.name}-{'直播' if payload.platform == 'live' else '短视频'}脚本"
        logger.info("Livestream script generated via AI service for '%s'", data["name"])
    else:
        logger.warning("AI service unavailable — falling back to local ai_service")
        content = ai_service.generate_live_script(data, payload.style, payload.platform)
        title = f"{product.name}-{'直播' if payload.platform == 'live' else '短视频'}脚本"

    script = LiveScript(
        product_id=product.id,
        title=title,
        style=payload.style,
        content=content,
    )
    db.add(script)
    db.commit()
    return AIResponse(result={"content": content, "title": title})


@router.post("/script/export")
def export_script(payload: dict, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    fmt = payload.get("format", "txt")
    content = payload.get("content", "")
    title = payload.get("title", "脚本")
    if fmt == "docx":
        from docx import Document
        import io
        doc = Document()
        doc.add_heading(title, level=1)
        for line in content.split("\n"):
            doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={title}.docx"})
    else:
        import io
        buffer = io.BytesIO(content.encode("utf-8"))
        from fastapi.responses import StreamingResponse
        return StreamingResponse(buffer, media_type="text/plain", headers={"Content-Disposition": f"attachment; filename={title}.txt"})


@router.get("/scripts")
def list_scripts(product_id: int = None, db: Session = Depends(get_db), _: User = Depends(require_merchant)):
    q = db.query(LiveScript)
    if product_id:
        q = q.filter(LiveScript.product_id == product_id)
    return q.order_by(LiveScript.created_at.desc()).all()
